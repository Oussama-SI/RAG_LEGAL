"""
نظام توليد العقود القانونية المغربية — Streamlit Community Cloud
Architecture: Kaggle RAG Server (FastAPI+ngrok) + Groq API (génération) → DOCX export
"""

import streamlit as st
from groq import Groq
import requests
import fitz  # PyMuPDF (upload PDF local optionnel)
import chromadb
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re
import io
import hashlib
import os

# ─── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="نظام العقود المغربية",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Amiri:wght@400;700&family=Cairo:wght@300;400;600;700&display=swap');

  html, body, [class*="css"] { direction: rtl; }

  .main { background: #f7f5f0; }

  h1, h2, h3 {
    font-family: 'Amiri', serif !important;
    color: #1a3a5c !important;
  }

  .stButton > button {
    font-family: 'Cairo', sans-serif !important;
    background: linear-gradient(135deg, #1a3a5c 0%, #2c5f8a 100%);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 0.6rem 1.5rem;
    font-size: 1rem;
    font-weight: 600;
    transition: all 0.2s ease;
    box-shadow: 0 2px 8px rgba(26,58,92,0.25);
  }
  .stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 16px rgba(26,58,92,0.35);
  }

  .contract-box {
    background: white;
    border: 1px solid #d0c8b8;
    border-radius: 12px;
    padding: 2rem;
    font-family: 'Amiri', serif;
    font-size: 1.1rem;
    line-height: 2;
    direction: rtl;
    text-align: right;
    box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    white-space: pre-wrap;
  }

  .metric-card {
    background: white;
    border-radius: 10px;
    padding: 1rem 1.5rem;
    text-align: center;
    border: 1px solid #e8e0d0;
    box-shadow: 0 2px 10px rgba(0,0,0,0.05);
  }

  .score-high { color: #1a7a3c; font-weight: 700; font-size: 1.4rem; }
  .score-med  { color: #c07020; font-weight: 700; font-size: 1.4rem; }
  .score-low  { color: #c02020; font-weight: 700; font-size: 1.4rem; }

  .sidebar-header {
    font-family: 'Amiri', serif;
    font-size: 1.3rem;
    font-weight: 700;
    color: #1a3a5c;
    border-bottom: 2px solid #c0a060;
    padding-bottom: 0.5rem;
    margin-bottom: 1rem;
  }

  .badge {
    display: inline-block;
    padding: 0.2rem 0.8rem;
    border-radius: 20px;
    font-size: 0.8rem;
    font-weight: 600;
    margin: 0.2rem;
  }
  .badge-green { background: #e8f5e9; color: #1a7a3c; }
  .badge-blue  { background: #e3eef7; color: #1a3a5c; }
  .badge-gold  { background: #fdf5e0; color: #8b6914; }
</style>
""", unsafe_allow_html=True)

# ─── Constants ────────────────────────────────────────────────────────────────
MOROCCAN_CONTRACT_INFO = {
    'عقد_إيجار': {
        'title': 'عقد كراء سكني',
        'law': 'القانون رقم 67.12 المتعلق بتنظيم العلاقات بين المكري والمكتري',
        'parties': ['المكري (الطرف الأول)', 'المكتري (الطرف الثاني)'],
        'clauses': ['وصف العقار', 'مدة الكراء', 'مبلغ الكراء وطريقة الأداء',
                    'الضمان', 'التزامات المكري', 'التزامات المكتري', 'فسخ العقد'],
        'icon': '🏠'
    },
    'عقد_بيع': {
        'title': 'عقد البيع',
        'law': 'ظهير الالتزامات والعقود — الفصول 478 إلى 618',
        'parties': ['البائع (الطرف الأول)', 'المشتري (الطرف الثاني)'],
        'clauses': ['وصف المبيع', 'الثمن وطريقة الأداء', 'نقل الملكية',
                    'ضمان الاستحقاق', 'التسليم', 'الفسخ'],
        'icon': '📜'
    },
    'عقد_عمل': {
        'title': 'عقد الشغل',
        'law': 'مدونة الشغل المغربية — القانون رقم 65.99',
        'parties': ['المشغل (الطرف الأول)', 'الأجير (الطرف الثاني)'],
        'clauses': ['طبيعة العمل', 'الأجر والامتيازات', 'مدة العقد',
                    'فترة التجربة', 'أوقات العمل', 'الإجازات', 'الإنهاء'],
        'icon': '💼'
    },
    'عقد_شراكة': {
        'title': 'عقد الشركة',
        'law': 'القانون رقم 5.96 المتعلق بشركات الأشخاص',
        'parties': ['الشريك الأول', 'الشريك الثاني'],
        'clauses': ['موضوع الشركة', 'رأس المال وتوزيع الحصص',
                    'توزيع الأرباح', 'تسيير الشركة', 'حل الشركة'],
        'icon': '🤝'
    },
    'عقد_مقاولة': {
        'title': 'عقد المقاولة',
        'law': 'ظهير الالتزامات والعقود — الفصول 723 إلى 769',
        'parties': ['صاحب المشروع (الطرف الأول)', 'المقاول (الطرف الثاني)'],
        'clauses': ['وصف الأشغال', 'الأثمان وطريقة الأداء', 'المدة', 'الضمانات'],
        'icon': '🏗️'
    },
    'عقد_قرض': {
        'title': 'عقد القرض',
        'law': 'ظهير الالتزامات والعقود — الفصول 860 إلى 877',
        'parties': ['المقرض (الطرف الأول)', 'المقترض (الطرف الثاني)'],
        'clauses': ['مبلغ القرض', 'الفائدة', 'مدة السداد', 'الضمانات', 'حالات الفسخ'],
        'icon': '💰'
    },
}

SYSTEM_PROMPT = """أنت محامٍ مغربي متخصص في صياغة العقود القانونية الرسمية وخبير في التشريع المغربي.
تصيغ عقوداً قانونية مغربية احترافية ومتوافقة مع القانون المغربي النافذ.

قواعد الصياغة الإلزامية:
١. ابدأ دائماً بـ "بسم الله الرحمن الرحيم" في سطر مستقل
٢. ثم "المملكة المغربية" في سطر مستقل
٣. أشر إلى القانون المغربي المنطبق في الديباجة
٤. حدد طرفي العقد بدقة مع بياناتهم الكاملة
٥. رقّم البنود بالأرقام العربية (البند الأول، البند الثاني...)
٦. استخدم المصطلحات القانونية المغربية الدقيقة
٧. اختم بخانة التوقيع: الطرف الأول والطرف الثاني والتاريخ والمكان
٨. أضف بند الفسخ وبند الاختصاص القضائي دائماً
٩. لا تضف أي تعليق أو شرح خارج نص العقد
١٠. اكتب العقد كاملاً بجميع بنوده دون اختصار"""

# ─── Arabic text utilities ────────────────────────────────────────────────────
def normalize_arabic(text: str) -> str:
    text = re.sub(r'ـ+', '', text)
    text = re.sub(r'[إأآا]', 'ا', text)
    text = re.sub(r'[\u064B-\u065F\u0670]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def clean_pdf_text(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    return '\n'.join(lines)

ARTICLE_RE = re.compile(
    r'(البند\s+(?:الأول|الثاني|الثالث|الرابع|الخامس|السادس|السابع|الثامن|التاسع|العاشر'
    r'|الحادي عشر|الثاني عشر|[\d\u0660-\u0669]+)'
    r'|المادة\s+(?:الأولى|الثانية|الثالثة|[\d\u0660-\u0669]+)'
    r'|الفصل\s+(?:الأول|الثاني|[\d\u0660-\u0669]+)'
    r'|أولاً|ثانياً|ثالثاً|رابعاً|خامساً)',
    re.UNICODE
)

def chunk_contract(text: str, max_size: int = 1200, min_size: int = 150) -> list:
    parts = ARTICLE_RE.split(text)
    chunks = []
    if len(parts) > 3:
        cur_art, cur_txt = 'مقدمة', ''
        for part in parts:
            if not part:
                continue
            if ARTICLE_RE.match(part):
                if len(cur_txt.strip()) >= min_size:
                    chunks.append({'text': cur_txt.strip(), 'article': cur_art})
                cur_art, cur_txt = part.strip(), part
            else:
                cur_txt += ' ' + part
                if len(cur_txt) > max_size:
                    chunks.append({'text': cur_txt[:max_size].strip(), 'article': cur_art})
                    cur_txt = cur_txt[max_size:]
        if len(cur_txt.strip()) >= min_size:
            chunks.append({'text': cur_txt.strip(), 'article': cur_art})
    else:
        step = max_size // 6
        words = text.split()
        for i in range(0, len(words), step - step // 5):
            chunk = ' '.join(words[i:i + step])
            if len(chunk) >= min_size:
                chunks.append({'text': chunk, 'article': f'قسم_{i // step + 1}'})
    return chunks

# ─── Session state init ───────────────────────────────────────────────────────
def init_state():
    defaults = {
        'chroma_client': None,
        'collection': None,
        'embed_model': None,
        'kaggle_api_url': '',
        'kaggle_status': {},
        'indexed_files': [],
        'last_contract': '',
        'last_contract_type': '',
        'last_metrics': {},
        'generation_history': [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ─── Kaggle RAG API helpers ───────────────────────────────────────────────────
def kaggle_health(api_url: str) -> dict:
    """Vérifie que le serveur Kaggle est actif."""
    try:
        r = requests.get(f"{api_url}/health", timeout=8)
        return r.json() if r.status_code == 200 else {}
    except Exception:
        return {}

def kaggle_retrieve(api_url: str, query: str, contract_type: str, n: int = 5) -> list:
    """Appelle /retrieve sur le serveur Kaggle."""
    try:
        r = requests.post(
            f"{api_url}/retrieve",
            json={"query": query, "contract_type": contract_type, "n": n},
            timeout=20
        )
        if r.status_code == 200:
            return r.json().get("chunks", [])
    except Exception as e:
        st.warning(f"RAG Kaggle non disponible: {e}")
    return []

def kaggle_generate(api_url: str, prompt: str) -> str | None:
    """Appelle /generate sur le serveur Kaggle (LLM local GPU)."""
    try:
        r = requests.post(
            f"{api_url}/generate",
            json={"prompt": prompt, "max_new_tokens": 2500},
            timeout=120
        )
        if r.status_code == 200:
            return r.json().get("contract", "")
    except Exception:
        return None

# ─── PDF parsing ──────────────────────────────────────────────────────────────
def parse_pdf_bytes(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text('text', flags=fitz.TEXT_PRESERVE_WHITESPACE)
        if text.strip():
            pages.append(f'[صفحة {i+1}]\n{text}')
    doc.close()
    return clean_pdf_text('\n\n'.join(pages))

def detect_contract_type(text: str) -> str:
    CONTRACT_KEYWORDS = {
        'عقد_إيجار':  ['إيجار', 'مستأجر', 'مؤجر', 'أجرة', 'كراء', 'مكتري', 'مكري'],
        'عقد_بيع':    ['بيع', 'مشتري', 'بائع', 'ثمن', 'ملكية', 'شراء'],
        'عقد_عمل':    ['عمل', 'عامل', 'صاحب العمل', 'راتب', 'أجر', 'توظيف', 'أجير', 'مشغل'],
        'عقد_شراكة':  ['شراكة', 'شريك', 'حصة', 'أرباح', 'خسائر', 'شركة'],
        'عقد_مقاولة': ['مقاولة', 'مقاول', 'أشغال', 'بناء', 'تشييد'],
        'عقد_قرض':    ['قرض', 'مقترض', 'مقرض', 'فائدة', 'دين', 'سلفة'],
    }
    sample = normalize_arabic(text[:3000])
    scores = {t: sum(kw in sample for kw in kws) for t, kws in CONTRACT_KEYWORDS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] >= 2 else 'عقد_عام'

# ─── Index PDFs : délégué au serveur Kaggle ──────────────────────────────────
# L'indexation se fait dans le notebook Kaggle (cellule 5).
# Streamlit ne fait qu'appeler l'API /retrieve.

# ─── RAG retrieval — délégué au serveur Kaggle ───────────────────────────────
def retrieve(query: str, contract_type: str = None, n: int = 5) -> list:
    api_url = st.session_state.get('kaggle_api_url', '').rstrip('/')
    if not api_url:
        return []
    return kaggle_retrieve(api_url, query, contract_type or '', n)

# ─── Build RAG prompt ─────────────────────────────────────────────────────────
def build_rag_prompt(request: str, contract_type: str, info: dict, context_chunks: list) -> str:
    ctx = ''
    if context_chunks:
        ctx = '\n\n## أمثلة مرجعية من عقود مغربية فعلية:\n'
        for i, c in enumerate(context_chunks[:4]):
            ctx += f'\n### مرجع {i+1} ({c["meta"].get("article", "")}):\n'
            ctx += c['text'][:600] + '...\n'

    clauses_list = '\n'.join(f'- {cl}' for cl in info.get('clauses', []))
    parties_list = '\n'.join(f'- {p}' for p in info.get('parties', []))

    return f"""## نوع العقد المطلوب: {info.get('title', contract_type)}
## القانون المغربي المنطبق: {info.get('law', 'القانون المدني المغربي')}

## أطراف العقد:
{parties_list}

## البنود الإلزامية الواجب تضمينها:
{clauses_list}

{ctx}

## متطلبات العقد المحددة:
{request}

اكتب العقد القانوني المغربي الكامل والمفصّل مع جميع البنود:"""

# ─── Generate contract via Groq API (gratuit) ────────────────────────────────
def generate_contract_groq(prompt: str, api_key: str) -> str:
    client = Groq(api_key=api_key)
    chat_completion = client.chat.completions.create(
        model="qwen-qwq-32b",   # Excellent arabe, gratuit sur Groq
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": prompt},
        ],
        max_tokens=4096,
        temperature=0.3,
        top_p=0.9,
    )
    return chat_completion.choices[0].message.content

# ─── Quality evaluation ───────────────────────────────────────────────────────
def evaluate_contract_quality(contract_text: str, contract_type: str) -> dict:
    info = MOROCCAN_CONTRACT_INFO.get(contract_type, {})
    required_clauses = info.get('clauses', [])

    m = {}
    m['bismillah']    = 'بسم الله' in contract_text
    m['maroc_header'] = 'المملكة المغربية' in contract_text
    m['has_parties']  = 'الطرف الأول' in contract_text and 'الطرف الثاني' in contract_text
    m['has_signature']= any(kw in contract_text for kw in ['التوقيع', 'توقيع', 'إمضاء'])
    m['has_date']     = any(kw in contract_text for kw in ['التاريخ', 'بتاريخ', '2025', '2024', '2026'])
    m['has_law_ref']  = any(kw in contract_text for kw in ['القانون', 'المادة', 'الفصل', 'ظهير'])
    m['has_jurisdiction'] = any(kw in contract_text for kw in ['الاختصاص', 'المحكمة', 'قضائي'])

    clause_found = []
    for clause in required_clauses:
        keywords = [kw for kw in clause.split() if len(kw) > 3]
        clause_found.append(any(kw in contract_text for kw in keywords))

    m['clause_coverage'] = sum(clause_found) / len(clause_found) if clause_found else 0
    m['clauses_found']   = sum(clause_found)
    m['clauses_total']   = len(clause_found)
    m['char_count']      = len(contract_text)
    m['word_count']      = len(contract_text.split())
    m['arabic_ratio']    = sum(1 for c in contract_text if '\u0600' <= c <= '\u06FF') / max(len(contract_text), 1)

    binary = [v for v in m.values() if isinstance(v, bool)]
    score = (sum(binary) / len(binary) * 0.4 + m['clause_coverage'] * 0.4 +
             min(m['arabic_ratio'] * 1.2, 1.0) * 0.2)
    m['overall_score'] = round(score * 100, 1)

    return m

# ─── DOCX export (enhanced quality) ──────────────────────────────────────────
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_arabic_font(run, size=12, bold=False, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Traditional Arabic')
    rPr.insert(0, rFonts)

def add_horizontal_line(doc, color=(26, 58, 92), thickness='12'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thickness)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '%02X%02X%02X' % tuple(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def contract_to_docx_bytes(contract_text: str, title: str, contract_type: str) -> bytes:
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.2)
    section.right_margin  = Cm(3.2)

    # ── Header: Kingdom of Morocco ──
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('المملكة المغربية')
    set_arabic_font(r, size=11, bold=True, color=(128, 0, 0))

    # ── Bismillah ──
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('بسم الله الرحمن الرحيم')
    set_arabic_font(r, size=15, bold=True)

    add_horizontal_line(doc, color=(192, 0, 0), thickness='18')

    # ── Contract title ──
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(title)
    set_arabic_font(r, size=20, bold=True, color=(26, 58, 92))

    add_horizontal_line(doc, color=(26, 58, 92), thickness='12')

    # ── Generation metadata ──
    info = MOROCCAN_CONTRACT_INFO.get(contract_type, {})
    if info.get('law'):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'استناداً إلى: {info["law"]}')
        set_arabic_font(r, size=10, italic=True, color=(100, 100, 100))

    doc.add_paragraph()  # spacing

    # ── Contract body ──
    lines = contract_text.split('\n')
    skip_headers = {'بسم الله الرحمن الرحيم', 'المملكة المغربية'}

    for line in lines:
        line = line.strip()
        if not line:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after  = Pt(2)
            continue

        # Skip duplicate headers already in our template
        if any(h in line for h in skip_headers) and lines.index(line) < 10:
            continue

        p = doc.add_paragraph()
        set_rtl(p)

        is_article = bool(re.match(
            r'^(البند|المادة|الفصل|أولاً|ثانياً|ثالثاً|رابعاً|خامساً)',
            line
        ))
        is_center_line = any(kw in line for kw in ['عقد ', 'بسم الله', 'المملكة'])
        is_signature_line = any(kw in line for kw in ['الطرف الأول', 'الطرف الثاني',
                                                        'الاسم:', 'التوقيع:', 'التاريخ:'])

        r = p.add_run(line)

        if is_article:
            set_arabic_font(r, size=13, bold=True, color=(26, 58, 92))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
        elif is_center_line:
            set_arabic_font(r, size=12, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_signature_line:
            set_arabic_font(r, size=11, bold=False, color=(60, 60, 60))
        else:
            set_arabic_font(r, size=11)

        p.paragraph_format.line_spacing = Pt(20)

    # ── Signature table ──
    doc.add_paragraph()
    add_horizontal_line(doc, color=(192, 160, 60), thickness='6')
    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'

    # Header row
    for i, cell in enumerate(sig_table.rows[0].cells):
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(['الطرف الأول', 'الطرف الثاني'][i])
        set_arabic_font(r, size=12, bold=True, color=(26, 58, 92))

    # Rows: name, date, signature
    labels = ['الاسم: ________________________________',
              'التاريخ: _____________________________',
              'التوقيع:\n\n\n']
    for row_i, label in enumerate(labels):
        for cell in sig_table.rows[row_i + 1].cells:
            p = cell.paragraphs[0]
            set_rtl(p)
            r = p.add_run(label)
            set_arabic_font(r, size=10)

    # ── Footer ──
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(
        f'تم التحرير بتاريخ: {datetime.now().strftime("%d/%m/%Y")}  |  '
        f'نموذج مولّد بالذكاء الاصطناعي — للاستخدام المرجعي فقط'
    )
    set_arabic_font(r, size=8, italic=True, color=(150, 150, 150))

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ═══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════════════════

# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown('<div class="sidebar-header">⚖️ نظام العقود المغربية</div>', unsafe_allow_html=True)

    # ── Groq API key ──
    st.markdown("**🔑 مفتاح Groq API (مجاني)**")
    api_key = st.text_input(
        "Groq Key", type="password", placeholder="gsk_...",
        label_visibility="collapsed"
    )
    if api_key:
        st.markdown('<span class="badge badge-green">✓ مفتاح Groq مُدخل</span>', unsafe_allow_html=True)

    st.divider()

    # ── Kaggle RAG server ──
    st.markdown("**🖥️ خادم RAG — Kaggle (ngrok)**")
    kaggle_url_input = st.text_input(
        "URL ngrok",
        value=st.session_state.get('kaggle_api_url', ''),
        placeholder="https://xxxx-xx-xx.ngrok-free.app",
        label_visibility="collapsed"
    )
    if kaggle_url_input != st.session_state.get('kaggle_api_url', ''):
        st.session_state.kaggle_api_url = kaggle_url_input.rstrip('/')
        st.session_state.kaggle_status = {}

    if st.button("🔌 Tester la connexion", use_container_width=True):
        url = st.session_state.get('kaggle_api_url', '')
        if url:
            with st.spinner("Vérification..."):
                st.session_state.kaggle_status = kaggle_health(url)
        else:
            st.warning("Entrez l'URL ngrok d'abord")

    status = st.session_state.get('kaggle_status', {})
    if status.get("status") == "ok":
        st.markdown(f'<span class="badge badge-green">✓ Connecté — {status.get("chunks_indexed",0)} chunks indexés</span>', unsafe_allow_html=True)
        if status.get("gpu_available"):
            st.markdown('<span class="badge badge-gold">⚡ GPU Kaggle actif</span>', unsafe_allow_html=True)
    elif status:
        st.error("✗ Serveur non joignable — vérifiez l'URL et le notebook Kaggle")
    else:
        st.caption("Non connecté → RAG désactivé, génération sans contexte")

    st.divider()

    # ── Generation mode ──
    st.markdown("**⚙️ Mode de génération**")
    use_kaggle_llm = st.toggle(
        "LLM Kaggle GPU (Qwen-7B)",
        value=False,
        help="Activé = génération via Qwen-7B sur Kaggle. Désactivé = Groq API cloud."
    )
    if use_kaggle_llm:
        st.markdown('<span class="badge badge-gold">🖥️ Qwen-7B — Kaggle GPU</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="badge badge-blue">☁️ Groq API — cloud</span>', unsafe_allow_html=True)

    st.divider()
    st.markdown("""
    <small style="color:#888; font-family:'Cairo',sans-serif; direction:rtl; display:block;">
    <b>Architecture:</b><br>
    📡 RAG ← Kaggle FastAPI + ngrok<br>
    🤖 LLM ← Groq (cloud) ou Kaggle GPU<br>
    📄 Export ← python-docx RTL
    </small>
    """, unsafe_allow_html=True)

# ─── Main content ─────────────────────────────────────────────────────────────
st.markdown("""
<h1 style="text-align:center; font-family:'Amiri',serif; font-size:2.2rem; margin-bottom:0.2rem;">
  نظام توليد العقود القانونية المغربية
</h1>
<p style="text-align:center; color:#666; font-family:'Cairo',sans-serif; margin-bottom:2rem;">
  مدعوم بالذكاء الاصطناعي • متوافق مع التشريع المغربي النافذ
</p>
""", unsafe_allow_html=True)

# ─── Form ──────────────────────────────────────────────────────────────────────
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown("**نوع العقد**")
    contract_type = st.selectbox(
        "نوع العقد",
        options=list(MOROCCAN_CONTRACT_INFO.keys()),
        format_func=lambda k: f"{MOROCCAN_CONTRACT_INFO[k]['icon']} {MOROCCAN_CONTRACT_INFO[k]['title']}",
        label_visibility="collapsed"
    )

    info = MOROCCAN_CONTRACT_INFO[contract_type]

    with st.expander("📋 البنود المطلوبة"):
        for cl in info['clauses']:
            st.write(f"• {cl}")

    with st.expander("⚖️ القانون المنطبق"):
        st.write(info['law'])

with col2:
    st.markdown("**متطلبات العقد** (اكتب التفاصيل بالعربية)")
    user_request = st.text_area(
        "متطلبات العقد",
        height=220,
        placeholder="""مثال: قم بصياغة عقد كراء سكني بالمواصفات التالية:
- المكري: السيد أحمد بنعلي، المقيم بالدار البيضاء...
- المكتري: السيد يوسف الأمراني...
- العقار: شقة من 3 غرف، الطابق الثاني...
- مبلغ الكراء: 4500 درهم شهرياً
- مدة الكراء: سنة واحدة من 1 يناير 2025""",
        label_visibility="collapsed"
    )

generate_col, _ = st.columns([1, 3])
with generate_col:
    generate_btn = st.button("✨ توليد العقد", type="primary", use_container_width=True)

# ─── Generation ───────────────────────────────────────────────────────────────
if generate_btn:
    if not user_request.strip():
        st.error("⚠️ الرجاء كتابة متطلبات العقد.")
    else:
        with st.spinner("🔍 استرجاع السياق من قاعدة البيانات..."):
            context = retrieve(user_request, contract_type, n=5)

        rag_info = f"تم استرجاع **{len(context)}** مقطع مرجعي"
        if context:
            best_score = context[0]['score']
            rag_info += f" (أفضل تشابه: {best_score:.0%})"
        st.info(rag_info)

        prompt = build_rag_prompt(user_request, contract_type, info, context)

        if use_kaggle_llm:
            spinner_msg = "⚙️ جاري توليد العقد عبر Kaggle GPU (Qwen-7B)..."
        else:
            spinner_msg = "⚙️ جاري توليد العقد عبر Groq API..."

        with st.spinner(spinner_msg):
            try:
                if use_kaggle_llm:
                    kaggle_url = st.session_state.get('kaggle_api_url', '')
                    if not kaggle_url:
                        st.error("⚠️ URL Kaggle non configurée — désactivez le toggle ou entrez l'URL.")
                        st.stop()
                    contract_text = kaggle_generate(kaggle_url, prompt)
                    if not contract_text:
                        st.error("Le serveur Kaggle n'a pas répondu. Vérifiez que le notebook tourne.")
                        st.stop()
                else:
                    if not api_key:
                        st.error("⚠️ الرجاء إدخال مفتاح Groq API في الشريط الجانبي.")
                        st.stop()
                    contract_text = generate_contract_groq(prompt, api_key)
                st.session_state.last_contract = contract_text
                st.session_state.last_contract_type = contract_type
                st.session_state.last_metrics = evaluate_contract_quality(contract_text, contract_type)
                st.session_state.generation_history.append({
                    'type': contract_type,
                    'title': info['title'],
                    'date': datetime.now().strftime('%H:%M:%S'),
                    'score': st.session_state.last_metrics['overall_score']
                })
                st.success("✅ تم توليد العقد بنجاح!")
            except Exception as e:
                st.error(f"خطأ في توليد العقد: {str(e)}")

# ─── Display result ───────────────────────────────────────────────────────────
if st.session_state.last_contract:
    st.divider()

    # Quality metrics
    m = st.session_state.last_metrics
    score = m.get('overall_score', 0)
    score_class = 'score-high' if score >= 75 else ('score-med' if score >= 50 else 'score-low')

    metrics_cols = st.columns(6)
    checks = [
        ('بسم الله', 'bismillah'),
        ('ترويسة المغرب', 'maroc_header'),
        ('الأطراف', 'has_parties'),
        ('التوقيع', 'has_signature'),
        ('تاريخ', 'has_date'),
        ('مرجع قانوني', 'has_law_ref'),
    ]
    for col, (label, key) in zip(metrics_cols, checks):
        val = m.get(key, False)
        icon = '✅' if val else '❌'
        col.markdown(f"""
        <div class="metric-card">
          <div style="font-size:1.4rem">{icon}</div>
          <div style="font-family:'Cairo',sans-serif; font-size:0.8rem; color:#555">{label}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown(f"""
    <div style="text-align:center; margin:1rem 0; font-family:'Cairo',sans-serif;">
      البنود المغطاة: <b>{m.get('clauses_found',0)}/{m.get('clauses_total',0)}</b>
      &nbsp;|&nbsp; عدد الكلمات: <b>{m.get('word_count',0):,}</b>
      &nbsp;|&nbsp; نسبة النص العربي: <b>{m.get('arabic_ratio',0)*100:.0f}%</b>
      &nbsp;|&nbsp; <span class="{score_class}">النقاط: {score}%</span>
    </div>
    """, unsafe_allow_html=True)

    # Contract display
    st.markdown("### 📄 نص العقد")
    st.markdown(
        f'<div class="contract-box">{st.session_state.last_contract}</div>',
        unsafe_allow_html=True
    )

    # Export actions
    st.markdown("")
    dl_col1, dl_col2 = st.columns([1, 3])

    with dl_col1:
        ctype = st.session_state.last_contract_type
        ctitle = MOROCCAN_CONTRACT_INFO.get(ctype, {}).get('title', 'عقد')

        docx_bytes = contract_to_docx_bytes(
            st.session_state.last_contract,
            ctitle,
            ctype
        )
        filename = f"{ctitle}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

        st.download_button(
            label="📥 تحميل DOCX",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with dl_col2:
        txt_bytes = st.session_state.last_contract.encode('utf-8')
        st.download_button(
            label="📋 تحميل نص (.txt)",
            data=txt_bytes,
            file_name=f"{ctitle}.txt",
            mime="text/plain",
            use_container_width=False
        )

# ─── History ──────────────────────────────────────────────────────────────────
if st.session_state.generation_history:
    st.divider()
    with st.expander(f"📊 سجل التوليد ({len(st.session_state.generation_history)} عقد)"):
        for entry in reversed(st.session_state.generation_history):
            score = entry['score']
            color = '#1a7a3c' if score >= 75 else ('#c07020' if score >= 50 else '#c02020')
            st.markdown(
                f"🕐 `{entry['date']}` — **{entry['title']}** — "
                f"<span style='color:{color}; font-weight:700'>{score}%</span>",
                unsafe_allow_html=True
            )
