"""
docx_export.py
=================
Export du contrat final en .docx (RTL, arabe). Code repris tel quel depuis
app_old.py (fonctions set_rtl / set_arabic_font / add_horizontal_line /
contract_to_docx_bytes), simplement déplacé dans son propre module pour que
app.py reste centré sur l'orchestration.
"""

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from contract_config import MOROCCAN_CONTRACT_INFO


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)


def set_arabic_font(run, size=12, bold=False, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0) if color is None else RGBColor(*color)

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Traditional Arabic')
    rPr.insert(0, rFonts)


def add_horizontal_line(doc, color=(26, 58, 92), thickness='12'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thickness)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '%02X%02X%02X' % tuple(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def contract_to_docx_bytes(contract_text: str, title: str, contract_type: str) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('المملكة المغربية')
    set_arabic_font(r, size=11, bold=True, color=(128, 0, 0))

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('بسم الله الرحمن الرحيم')
    set_arabic_font(r, size=15, bold=True)

    add_horizontal_line(doc, color=(192, 0, 0), thickness='18')

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_arabic_font(r, size=20, bold=True, color=(26, 58, 92))

    add_horizontal_line(doc, color=(26, 58, 92), thickness='12')

    info = MOROCCAN_CONTRACT_INFO.get(contract_type, {})
    if info.get('law'):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'استناداً إلى: {info["law"]}')
        set_arabic_font(r, size=10, italic=True, color=(100, 100, 100))

    doc.add_paragraph()

    skip_headers = ['بسم الله الرحمن الرحيم', 'المملكة المغربية']
    header_count = {'بسم الله الرحمن الرحيم': 0, 'المملكة المغربية': 0}

    for line in contract_text.split('\n'):
        line = line.strip()
        if not line:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after = Pt(2)
            continue

        skip = False
        for h in skip_headers:
            if h in line:
                header_count[h] += 1
                if header_count[h] > 1:
                    skip = True
                break
        if skip:
            continue

        p = doc.add_paragraph()
        set_rtl(p)

        is_article = bool(re.match(r'^(البند|المادة|الفصل|أولاً|ثانياً|ثالثاً|رابعاً|خامساً)', line))
        is_center_line = any(kw in line for kw in ['عقد ', 'بسم الله', 'المملكة'])
        is_signature_line = any(kw in line for kw in ['الطرف الأول', 'الطرف الثاني', 'الاسم:', 'التوقيع:', 'التاريخ:'])

        r = p.add_run(line)

        if is_article:
            set_arabic_font(r, size=13, bold=True, color=(26, 58, 92))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif is_center_line:
            set_arabic_font(r, size=12, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_signature_line:
            set_arabic_font(r, size=11, bold=False, color=(60, 60, 60))
        else:
            set_arabic_font(r, size=11)

        p.paragraph_format.line_spacing = Pt(20)

    doc.add_paragraph()
    add_horizontal_line(doc, color=(192, 160, 60), thickness='6')
    doc.add_paragraph()

    # sig_table = doc.add_table(rows=4, cols=2)
    # sig_table.style = 'Table Grid'

    # for i, cell in enumerate(sig_table.rows[0].cells):
    #     p = cell.paragraphs[0]
    #     set_rtl(p)
    #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     r = p.add_run(['الطرف الأول', 'الطرف الثاني'][i])
    #     set_arabic_font(r, size=12, bold=True, color=(26, 58, 92))

    # labels = ['الاسم: ________________________________',
    #           'التاريخ: _____________________________',
    #           'التوقيع:\n\n\n']
    # for row_i, label in enumerate(labels):
    #     for cell in sig_table.rows[row_i + 1].cells:
    #         p = cell.paragraphs[0]
    #         set_rtl(p)
    #         r = p.add_run(label)
    #         set_arabic_font(r, size=10)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(
        f'تم التحرير بتاريخ: {datetime.now().strftime("%d/%m/%Y")}  |  '
        f'نموذج مولّد بالذكاء الاصطناعي — للاستخدام المرجعي فقط'
    )
    set_arabic_font(r, size=8, italic=True, color=(150, 150, 150))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
